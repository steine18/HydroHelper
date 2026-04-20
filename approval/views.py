import io
import json
import logging

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.views.decorators.http import require_POST
from docx import Document
from docx.shared import Pt, RGBColor

from sites.models import Site
from water_balance.usgs import USGSAPIError

from .approval_types import APPROVAL_TYPES_BY_ID
from .forms import NewApprovalForm
from .models import ApprovalRequest

logger = logging.getLogger(__name__)


@login_required
def index(request):
    approvals = ApprovalRequest.objects.filter(user=request.user).select_related('site')
    drafts = [a for a in approvals if a.status == 'draft']
    completed = [a for a in approvals if a.status == 'complete']
    return render(request, 'approval/index.html', {'drafts': drafts, 'completed': completed})


@login_required
def new_approval(request):
    if request.method == 'POST':
        form = NewApprovalForm(request.POST)
        if form.is_valid():
            site_no = form.cleaned_data['site_no']
            try:
                site = Site.get_or_fetch(site_no)
            except USGSAPIError as exc:
                form.add_error('site_no', str(exc))
                return render(request, 'approval/new_approval.html', {'form': form})
            approval = ApprovalRequest.objects.create(
                user=request.user,
                site=site,
                approval_type=form.cleaned_data['approval_type'],
                period_start=form.cleaned_data['period_start'],
                period_end=form.cleaned_data['period_end'],
            )
            return redirect('approval_detail', pk=approval.pk)
    else:
        form = NewApprovalForm()
    return render(request, 'approval/new_approval.html', {'form': form})


@login_required
def approval_detail(request, pk):
    approval = get_object_or_404(ApprovalRequest, pk=pk, user=request.user)
    at = APPROVAL_TYPES_BY_ID.get(approval.approval_type, {})
    items = at.get('items', [])
    return render(request, 'approval/approval_detail.html', {
        'approval': approval,
        'items': items,
        'completion_pct': approval.completion_pct(),
        'response_json': json.dumps(approval.response_data),
    })


@login_required
@require_POST
def autosave(request, pk):
    approval = get_object_or_404(ApprovalRequest, pk=pk, user=request.user)
    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, AttributeError):
        return JsonResponse({'ok': False}, status=400)
    approval.response_data = data
    try:
        approval.save(update_fields=['response_data', 'updated_at'])
    except Exception:
        logger.exception('autosave failed for ApprovalRequest pk=%s', pk)
        return JsonResponse({'ok': False}, status=500)
    return JsonResponse({'ok': True})


@login_required
@require_POST
def toggle_complete(request, pk):
    approval = get_object_or_404(ApprovalRequest, pk=pk, user=request.user)
    approval.status = 'complete' if approval.status == 'draft' else 'draft'
    approval.save(update_fields=['status', 'updated_at'])
    return redirect('approval_detail', pk=approval.pk)


@login_required
@require_POST
def update_dates(request, pk):
    approval = get_object_or_404(ApprovalRequest, pk=pk, user=request.user)
    period_start = request.POST.get('period_start', '').strip()
    period_end = request.POST.get('period_end', '').strip()
    if period_start and period_end:
        approval.period_start = period_start
        approval.period_end = period_end
        approval.save(update_fields=['period_start', 'period_end', 'updated_at'])
    return redirect('approval_detail', pk=approval.pk)


@login_required
@require_POST
def delete_approval(request, pk):
    approval = get_object_or_404(ApprovalRequest, pk=pk, user=request.user)
    approval.delete()
    return redirect('approval_index')


@login_required
def approval_report(request, pk):
    approval = get_object_or_404(ApprovalRequest, pk=pk, user=request.user)
    at = APPROVAL_TYPES_BY_ID.get(approval.approval_type, {})
    items = at.get('items', [])
    response_data = approval.response_data

    # Determine which questions are visible using the same conditional logic as the frontend
    visible = {}
    for item in items:
        if item['type'] == 'section':
            continue
        key = item['key']
        cond_on = item.get('conditional_on')
        if not cond_on:
            visible[key] = True
        else:
            cond_val = item.get('conditional_value', 'yes')
            parent_visible = visible.get(cond_on, False)
            parent_answer = response_data.get(cond_on, {}).get('answer', '')
            visible[key] = parent_visible and (parent_answer == cond_val)

    # Build report rows, skipping hidden conditional questions
    report_items = []
    for item in items:
        if item['type'] == 'section':
            report_items.append(item)
            continue
        key = item['key']
        if not visible.get(key, True):
            continue
        response = response_data.get(key, {})
        if item['type'] == 'yn':
            answer = response.get('answer', '')
            ans_label = ''
            ans_class = 'answer-blank'
            for opt in item.get('options', []):
                if opt['value'] == answer:
                    ans_label = opt['label']
                    if opt['value'] == 'na':
                        ans_class = 'answer-na'
                    elif opt['good']:
                        ans_class = 'answer-yes'
                    else:
                        ans_class = 'answer-no'
                    break
            response = {**response, 'ans_class': ans_class, 'ans_label': ans_label}
        report_items.append({**item, 'response': response})

    return render(request, 'approval/report.html', {
        'approval': approval,
        'report_items': report_items,
    })


# RGB color constants for docx
_COLOR_GREEN  = RGBColor(0x19, 0x87, 0x54)
_COLOR_RED    = RGBColor(0xDC, 0x35, 0x45)
_COLOR_GRAY   = RGBColor(0x6C, 0x75, 0x7D)
_COLOR_ORANGE = RGBColor(0xFD, 0x7E, 0x14)
_COLOR_LIGHT  = RGBColor(0xAD, 0xB5, 0xBD)


def _add_colored_run(para, text, rgb, bold=False):
    run = para.add_run(text)
    run.font.color.rgb = rgb
    run.bold = bold


@login_required
def export_docx(request, pk):
    approval = get_object_or_404(ApprovalRequest, pk=pk, user=request.user)
    at = APPROVAL_TYPES_BY_ID.get(approval.approval_type, {})
    items = at.get('items', [])
    response_data = approval.response_data

    # Same visibility logic as approval_report
    visible = {}
    for item in items:
        if item['type'] == 'section':
            continue
        key = item['key']
        cond_on = item.get('conditional_on')
        if not cond_on:
            visible[key] = True
        else:
            cond_val = item.get('conditional_value', 'yes')
            parent_visible = visible.get(cond_on, False)
            parent_answer = response_data.get(cond_on, {}).get('answer', '')
            visible[key] = parent_visible and (parent_answer == cond_val)

    doc = Document()

    # Document heading
    heading = doc.add_paragraph()
    heading.add_run(f"{approval.site.site_no} — {approval.site.name}").bold = True
    doc.add_paragraph(
        f"{approval.get_approval_type_display()} Approval  ·  "
        f"{approval.period_start} to {approval.period_end}  ·  "
        f"{'Complete' if approval.status == 'complete' else 'Draft'}"
    ).runs[0].font.color.rgb = _COLOR_GRAY

    for item in items:
        if item['type'] == 'section':
            doc.add_paragraph()
            sec = doc.add_paragraph()
            run = sec.add_run(f"{item['number']}. {item['title']}")
            run.bold = True
            run.font.size = Pt(11)
            continue

        key = item['key']
        if not visible.get(key, True):
            continue

        response = response_data.get(key, {})
        para = doc.add_paragraph()
        # Question number + text
        para.add_run(f"{item['number']}  ").font.color.rgb = _COLOR_GRAY
        para.add_run(item['text'])

        if item['type'] == 'yn':
            answer = response.get('answer', '')
            comment = response.get('comment', '')
            ans_label = ''
            ans_color = _COLOR_LIGHT
            for opt in item.get('options', []):
                if opt['value'] == answer:
                    ans_label = opt['label']
                    if opt['value'] == 'na':
                        ans_color = _COLOR_GRAY
                    elif opt['good']:
                        ans_color = _COLOR_GREEN
                    else:
                        ans_color = _COLOR_RED
                    break
            ans_para = doc.add_paragraph(style='List Bullet')
            if ans_label:
                _add_colored_run(ans_para, ans_label, ans_color, bold=True)
                if comment:
                    ans_para.add_run(f": {comment}")
            else:
                _add_colored_run(ans_para, '— not answered', _COLOR_LIGHT)
                if comment:
                    ans_para.add_run(f", {comment}")

        elif item['type'] == 'date':
            date_val = response.get('date', '')
            comment = response.get('comment', '')
            date_para = doc.add_paragraph(style='List Bullet')
            if date_val:
                date_para.add_run(date_val).bold = True
                if comment:
                    date_para.add_run(f"  —  {comment}")
            else:
                _add_colored_run(date_para, '— not answered', _COLOR_LIGHT)
                if comment:
                    date_para.add_run(f", {comment}")

        elif item['type'] == 'text':
            text_val = response.get('text', '')
            color_val = response.get('color', '')
            text_para = doc.add_paragraph()
            if text_val:
                color_map = {'green': _COLOR_GREEN, 'orange': _COLOR_ORANGE, 'red': _COLOR_RED}
                run = text_para.add_run(text_val)
                if color_val in color_map:
                    run.font.color.rgb = color_map[color_val]
            else:
                _add_colored_run(text_para, '— not answered', _COLOR_LIGHT)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    site_no = approval.site.site_no
    filename = f"approval_{site_no}_{approval.period_start}_{approval.period_end}.docx"
    response = HttpResponse(
        buf.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
